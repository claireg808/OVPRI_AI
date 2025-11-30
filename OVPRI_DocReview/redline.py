## Produce a revised "redlined" version of a user's provided document

import io
import os
import re
import magic
import docx2txt
from docx import Document
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from pdfminer.high_level import extract_text

load_dotenv()

# initialize LLM
llm = ChatOpenAI(
    base_url=os.environ['BASE_URL'],
    api_key='dummy-key',
    model=os.environ['MODEL'],
    temperature=0.2
)

# read checklist
with open('/home/gillaspiecl/OVPRI_AI/OVPRI_DocReview/data/CDA_Checklist.txt', 'r', encoding='utf-8') as f:
    checklist = f.read()

# read policy
with open('/home/gillaspiecl/OVPRI_AI/OVPRI_DocReview/data/Corporate_Research_Agreements_Policy.txt', 'r', encoding='utf-8') as f:
    policy = f.read()

# read redline prompt
with open('/home/gillaspiecl/OVPRI_AI/OVPRI_DocReview/doc_review_prompt.txt', 'r', encoding='utf-8') as f:
    redline_prompt_txt = f.read()

# read redline prompt
with open('/home/gillaspiecl/OVPRI_AI/OVPRI_DocReview/second_pass_prompt.txt', 'r', encoding='utf-8') as f:
    revision_prompt_txt = f.read()

# split into chunks
def split(text):
    # normalize and split
    text = re.sub(r'Page\s*\d.*\d', ' ', text)
    text = re.sub(r'\s*\n\s*', r'\n\n', text)
    chunks = re.split(r'(?=\b\d+\.)', text)

    # keep to 10 chunks or less - reduce # of llm calls
    while len(chunks) > 10:
        combined_pairs = []
        for i in range(0, len(chunks), 2):
            if i + 1 < len(chunks):
                # combine pairs into single chunk
                combined_pairs.append(chunks[i] + chunks[i+1])
            else:
                # odd number of chunks
                combined_pairs.append(chunks[i]) 
        chunks = combined_pairs

    return chunks

# format as doc
def format(text):
    document = Document()

    title = text.split('\n')[0]

    document.add_heading(title, level=1)

    body_text = text.split('\n')[1:]

    for paragraph in body_text:
        document.add_paragraph(paragraph)

    return document


# produce redline version of uploaded document
def redline_document(document, type):
    # read document type and reset file pointer
    file_type = magic.from_buffer(document.read(2048))
    document.seek(0)

    if 'pdf' in file_type.lower():  # .pdf
        pdf_bytes = io.BytesIO(document.read())
        text = extract_text(pdf_bytes)
    else:  # .docx
        try:
            text = docx2txt.process(document)
        except:
            return 'Please upload in either .pdf or .docx format.', {'error': 'Incorrect file type'}

    text_chunks = split(text)

    full_redline = ''
    for chunk in text_chunks:
        # generate redline prompt
        chunk_word_count = str(int(len(chunk) * 1.1))
        if type == 'Confidentiality Agreement':
            redline_prompt = redline_prompt_txt \
                        .replace('{POLICIES}', policy) \
                        .replace('{CHECKLIST}', checklist) \
                        .replace('{DOCUMENT}', text) \
                        .replace('{WORDS}', chunk_word_count)
        
            # invoke LLM for redlined section
            chunk_redline = llm.invoke(redline_prompt).content

            full_redline += chunk_redline

    full_word_count = str(int(len(text) * 1.2))
    revision_prompt = revision_prompt_txt \
                .replace('{POLICIES}', policy) \
                .replace('{CHECKLIST}', checklist) \
                .replace('{DOCUMENT}', full_redline) \
                .replace('{WORDS}', full_word_count)
    
    final_review = llm.invoke(revision_prompt).content

    redline_doc = format(final_review)
    final_doc = format(final_review)

    redline_doc.save('/home/gillaspiecl/OVPRI_AI/OVPRI_DocReview/data/AI_Redline_Test.docx')
    final_doc.save('/home/gillaspiecl/OVPRI_AI/OVPRI_DocReview/data/AI_Redline_Final_Test.docx')

    return final_review